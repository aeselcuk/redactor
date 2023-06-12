import re
import logging

from docx import Document

logger = logging.getLogger("docx-redact")

class DocxRedactor:
    def __init__(self, doc_obj_path, regexes, replace_char):
        self.doc_obj_path = doc_obj_path
        self.regexes = regexes
        self.replace_char = replace_char
        
    def __replace_text__(self, text):
        """ Helper function replaces all the characters
        with the replacment char except the empty spaces. """
        return re.sub(r"[^\s-]", lambda m: self.replace_char * len(m.group(0)), text)
        
    def __redact_helper__(self, doc_obj):
        """
        Helper function for the redact function
        :param doc_obj (Document): the whole document as in the input .docx file
        :return:
        """
        for reg in self.regexes:
            regex = re.compile(reg)
            for p in doc_obj.paragraphs:
                if regex.search(p.text):
                    inline = p.runs
                    for i in range(len(inline)):
                        if regex.search(inline[i].text):
                            text = re.sub(reg, lambda m: self.__replace_text__(m.group(0)), inline[i].text)
                            inline[i].text = text
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.__redact_helper__(cell)

    def redact(self, output_file_path):
        """
        Redacts the given .docx file and writes result to output file
        :param output_file_path (string): path of the file to write the result to
        :return:
        """
        doc_obj = Document(self.doc_obj_path)
        self.__redact_helper__(doc_obj)
        doc_obj.save(output_file_path)
        if self.doc_obj_path == output_file_path:
            logger.warning("Input and Output files are same!")
        logging.info("Updated file saved as: " + output_file_path)

