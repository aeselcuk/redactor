import tkinter as tk
from tkinter.filedialog import askopenfilename
from tkinter import Entry
from tkinter import messagebox
from tkinter import ttk
from docx import *
from docx.enum.text import WD_COLOR_INDEX
import os
import re

window = tk.Tk()
window.title("File Redaction")
window.geometry("500x300")
suffix = ""

def choose_file():
    global filename
    filename.set(askopenfilename())

def redact_file():
    global suffix
    suffix = suffix_entry.get()
    document = Document(filename.get())
    for p in document.paragraphs:
        start = p.text.find('/')
        end = p.text.find('/', start+1) 
        if start > -1 and end > -1:
            redactedPart = re.sub('[a-zA-Z0-9]', "x", p.text[start+1:end])
            pre = p.text[:start]
            post = p.text[end+1:]
            p.text = pre
            p.add_run(redactedPart)
            p.runs[1].font.highlight_color = WD_COLOR_INDEX.BLACK
            p.add_run(post)
    filepath, fileExtension = os.path.splitext(filename.get())
    fullredactedpath = filepath + suffix + fileExtension
    document.save(fullredactedpath)
    if os.path.isfile(fullredactedpath):
        messagebox.showinfo(title=None, message="Successfully Redacted at "+fullredactedpath)



appname_display = tk.Label(window, text="File Redactor", font=("Calibri 20 bold"))
appname_display.place(x=250, y=30, anchor="center")

step1_display = tk.Label(window, text="1- Choose file to be redacted", font=("Calibri 12 bold"))
step1_display.place(x=20, y=80, anchor="w")

choose_file_button = tk.Button(window,text="Choose File", font=("Calibri 12"),command=choose_file)
choose_file_button.place(x=20, y=120, anchor="w")

filename = tk.StringVar(window, "No file selected.")
filename_display = tk.Label(window,textvariable=filename,font=("Calibri 11"))
filename_display.place(x=110,y=120, anchor="w")

step2_display = tk.Label(window, text="2- Add suffix for the redacted file", font=("Calibri 12 bold"))
step2_display.place(x=20, y=160, anchor="w")

suffix_entry = Entry(window, width=30)
suffix_entry.pack()
suffix_entry.place(x=20, y=190, anchor="w")

step3_display = tk.Label(window, text="3- Click below to redact your file", font=("Calibri 12 bold"))
step3_display.place(x=20, y=230, anchor="w")

redact_file_button = tk.Button(window,text="Redact File", font=("Calibri 12 bold"), command=redact_file)
redact_file_button.place(x=20,y=270, anchor="w")

window.mainloop()